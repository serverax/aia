from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.editor_agent.services.models import Template, TemplateSection


class DocxGenerator:
    def __init__(self, template: Template):
        self.template = template
        self.doc = Document()

    def generate(self, content: dict) -> bytes:
        for section in self.template.sections:
            value = content.get(section.field, "")

            if section.type == "title":
                paragraph = self.doc.add_paragraph(str(value))
                paragraph.style = "Heading 1"
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif section.type == "text":
                self.doc.add_paragraph(str(value))

            elif section.type == "table":
                self._add_table(section, value if isinstance(value, list) else [])

            elif section.type == "references":
                self._add_references(value if isinstance(value, list) else [])

        buffer = io.BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()

    def _add_table(self, section: TemplateSection, rows: list[dict]):
        columns = section.columns or []
        if not columns:
            return
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(columns))

        for index, column_name in enumerate(columns):
            table.rows[0].cells[index].text = column_name

        for row_index, row_data in enumerate(rows, start=1):
            for column_index, column_name in enumerate(columns):
                table.rows[row_index].cells[column_index].text = str(row_data.get(column_name, ""))

    def _add_references(self, policies: list[dict]):
        self.doc.add_paragraph("References:", style="Heading 3")
        for policy in policies:
            title = policy.get("title", "Untitled policy")
            jurisdiction = policy.get("jurisdiction", "N/A")
            self.doc.add_paragraph(f"{title} ({jurisdiction})", style="List Bullet")

