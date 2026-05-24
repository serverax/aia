import markdown
from typing import List, Dict, Any, Optional
from uuid import UUID
from .models.finalizer import DocumentDraft, FinalDocument, DocumentFormat
import hashlib
import time
import os
from docx import Document
from fpdf import FPDF


class EditorAgent:
    def __init__(self, reports_dir="F:/aia/reports"):
        self.reports_dir = reports_dir
        if not os.path.exists(self.reports_dir):
            os.makedirs(self.reports_dir)

        # In-memory stores for MVP
        self.drafts: Dict[UUID, DocumentDraft] = {}
        self.finalized: Dict[UUID, FinalDocument] = {}

    def create_draft(self, project_id: str, agent_id: str, content: str) -> DocumentDraft:
        """Convert LLM markdown to versioned draft with HTML preview."""
        html_content = markdown.markdown(content)

        existing_versions = [d for d in self.drafts.values() if d.project_id == project_id]
        version = len(existing_versions) + 1

        draft = DocumentDraft(
            project_id=project_id,
            agent_id=agent_id,
            raw_content=content,
            formatted_content=html_content,
            version=version,
        )
        self.drafts[draft.draft_id] = draft
        return draft

    def get_preview(self, project_id: str, version: Optional[int] = None) -> str:
        """Get HTML preview for a project version."""
        project_drafts = [d for d in self.drafts.values() if d.project_id == project_id]
        if not project_drafts:
            return "<p>No drafts found for this project.</p>"

        if version:
            draft = next((d for d in project_drafts if d.version == version), None)
        else:
            draft = max(project_drafts, key=lambda d: d.version)

        return draft.formatted_content if draft else "<p>Version not found.</p>"

    def finalize_document(self, project_id: str, format: DocumentFormat) -> FinalDocument:
        """Finalize drafts into a real binary file (DOCX/PDF)."""
        project_drafts = [d for d in self.drafts.values() if d.project_id == project_id]
        if not project_drafts:
            raise ValueError(f"No drafts found to finalize for project {project_id}")

        latest_draft = max(project_drafts, key=lambda d: d.version)
        filename = f"{project_id}_{int(time.time())}.{format.lower()}"
        file_path = os.path.join(self.reports_dir, filename)

        if format == DocumentFormat.DOCX:
            self._generate_docx(latest_draft.raw_content, file_path, project_id)
        elif format == DocumentFormat.PDF:
            self._generate_pdf(latest_draft.raw_content, file_path, project_id)

        signature = hashlib.sha256(latest_draft.raw_content.encode()).hexdigest()

        final_doc = FinalDocument(
            project_id=project_id,
            format=format,
            file_url=f"/reports/{filename}",  # Local relative path for serving
            signature_hash=signature,
            audit_trail=[d.draft_id for d in project_drafts],
        )
        self.finalized[final_doc.document_id] = final_doc
        return final_doc

    def _generate_docx(self, content: str, path: str, project_id: str):
        doc = Document()
        doc.add_heading(f"Compliance Report: {project_id}", 0)
        doc.add_paragraph(content)
        doc.save(path)

    def _generate_pdf(self, content: str, path: str, project_id: str):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(40, 10, f"Compliance Report: {project_id}")
        pdf.ln(20)
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 10, content)
        pdf.output(path)

    def get_audit_history(self, document_id: UUID) -> List[DocumentDraft]:
        """Retrieve the sequence of drafts for a finalized document."""
        doc = self.finalized.get(document_id)
        if not doc:
            return []
        return [self.drafts[draft_id] for draft_id in doc.audit_trail if draft_id in self.drafts]
